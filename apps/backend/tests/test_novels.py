"""Novel / Chapter 测试（Phase 2 — Novel Studio）。"""

import io


def _create_project(client) -> str:
    return client.post("/api/projects", json={"name": "小说项目"}).json()["id"]


def test_novel_crud(client):
    pid = _create_project(client)
    created = client.post(f"/api/projects/{pid}/novels", json={"title": "我的小说"})
    assert created.status_code == 201
    novel_id = created.json()["id"]

    listing = client.get(f"/api/projects/{pid}/novels")
    assert listing.status_code == 200
    assert any(n["id"] == novel_id for n in listing.json())

    detail = client.get(f"/api/projects/{pid}/novels/{novel_id}")
    assert detail.status_code == 200
    assert detail.json()["novel"]["title"] == "我的小说"
    assert detail.json()["chapters"] == []

    updated = client.put(
        f"/api/projects/{pid}/novels/{novel_id}", json={"title": "改名"}
    )
    assert updated.status_code == 200
    assert updated.json()["title"] == "改名"

    deleted = client.delete(f"/api/projects/{pid}/novels/{novel_id}")
    assert deleted.status_code == 204
    assert client.get(f"/api/projects/{pid}/novels/{novel_id}").status_code == 404


def test_create_novel_empty_title_clear_message(client):
    pid = _create_project(client)
    empty = client.post(f"/api/projects/{pid}/novels", json={"title": ""})
    assert empty.status_code == 422
    body = empty.json()["error"]
    assert body["code"] == "validation_error"
    assert "标题" in body["message"]
    assert "不能为空" in body["message"]

    missing = client.post(f"/api/projects/{pid}/novels", json={})
    assert missing.status_code == 422
    body = missing.json()["error"]
    assert body["code"] == "validation_error"
    assert "缺少参数" in body["message"]
    assert "标题" in body["message"]


def test_chapter_crud(client):
    pid = _create_project(client)
    novel_id = client.post(
        f"/api/projects/{pid}/novels", json={"title": "章节测试"}
    ).json()["id"]
    added = client.post(
        f"/api/projects/{pid}/novels/{novel_id}/chapters", json={"title": "第一章"}
    )
    assert added.status_code == 201
    chapter_id = added.json()["id"]

    updated = client.put(
        f"/api/projects/{pid}/novels/{novel_id}/chapters/{chapter_id}",
        json={"content": "正文内容"},
    )
    assert updated.status_code == 200
    assert updated.json()["content"] == "正文内容"

    detail = client.get(f"/api/projects/{pid}/novels/{novel_id}").json()
    assert len(detail["chapters"]) == 1
    assert detail["chapters"][0]["order_index"] == 0

    deleted = client.delete(
        f"/api/projects/{pid}/novels/{novel_id}/chapters/{chapter_id}"
    )
    assert deleted.status_code == 204
    assert client.get(f"/api/projects/{pid}/novels/{novel_id}").json()["chapters"] == []


def test_search_novels(client):
    pid = _create_project(client)
    n1 = client.post(f"/api/projects/{pid}/novels", json={"title": "修仙传"}).json()["id"]
    n2 = client.post(f"/api/projects/{pid}/novels", json={"title": "都市"}).json()["id"]
    chapter = client.post(
        f"/api/projects/{pid}/novels/{n1}/chapters", json={"title": "第一章"}
    ).json()
    client.put(
        f"/api/projects/{pid}/novels/{n1}/chapters/{chapter['id']}",
        json={"content": "主角林风在青云门修炼"},
    )

    found = client.get(f"/api/projects/{pid}/novels", params={"q": "青云门"}).json()
    assert [n["id"] for n in found] == [n1]
    found2 = client.get(f"/api/projects/{pid}/novels", params={"q": "都市"}).json()
    assert [n["id"] for n in found2] == [n2]


def test_import_markdown_splits_chapters(client):
    pid = _create_project(client)
    content = (
        "# 第一章 出山\n正文一\n\n## 第一节\n正文二\n# 第二章 下山\n正文三"
    )
    response = client.post(
        f"/api/projects/{pid}/novels/import?filename=novel.md",
        content=content.encode("utf-8"),
        headers={"Content-Type": "text/markdown"},
    )
    assert response.status_code == 201
    novel = response.json()
    assert novel["source_type"] == "imported"
    detail = client.get(f"/api/projects/{pid}/novels/{novel['id']}").json()
    titles = [c["title"] for c in detail["chapters"]]
    assert titles == ["第一章 出山", "第一节", "第二章 下山"]
    assert "正文一" in detail["chapters"][0]["content"]
    assert detail["chapters"][1]["content"] == "正文二"
    assert "正文三" in detail["chapters"][2]["content"]


def test_import_gbk_txt(client):
    pid = _create_project(client)
    content = "第一章 测试\n这是GBK编码的中文内容".encode("gbk")
    response = client.post(
        f"/api/projects/{pid}/novels/import?filename=gbk.txt",
        content=content,
        headers={"Content-Type": "text/plain"},
    )
    assert response.status_code == 201
    detail = client.get(f"/api/projects/{pid}/novels/{response.json()['id']}").json()
    assert "这是GBK编码的中文内容" in detail["chapters"][0]["content"]


def test_import_docx_splits_chapters(client):
    from docx import Document

    pid = _create_project(client)
    doc = Document()
    doc.add_heading("第一章 风起", level=1)
    doc.add_paragraph("山门外，风雪正紧。")
    doc.add_heading("第二章 云涌", level=1)
    doc.add_paragraph("林风推门而入。")
    buffer = io.BytesIO()
    doc.save(buffer)

    response = client.post(
        f"/api/projects/{pid}/novels/import?filename=novel.docx",
        content=buffer.getvalue(),
        headers={
            "Content-Type": (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        },
    )
    assert response.status_code == 201
    detail = client.get(f"/api/projects/{pid}/novels/{response.json()['id']}").json()
    assert [c["title"] for c in detail["chapters"]] == ["第一章 风起", "第二章 云涌"]
    assert "山门外" in detail["chapters"][0]["content"]


def test_import_unsupported_format(client):
    pid = _create_project(client)
    response = client.post(
        f"/api/projects/{pid}/novels/import?filename=data.pdf",
        content=b"x",
        headers={"Content-Type": "application/pdf"},
    )
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "unsupported_format"


def test_import_missing_filename(client):
    pid = _create_project(client)
    response = client.post(
        f"/api/projects/{pid}/novels/import",
        content=b"x",
        headers={"Content-Type": "text/plain"},
    )
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "validation_error"


def test_novel_project_not_found(client):
    response = client.post("/api/projects/proj_missing/novels", json={"title": "x"})
    assert response.status_code == 404
    assert response.json()["error"]["code"] == "project_not_found"
